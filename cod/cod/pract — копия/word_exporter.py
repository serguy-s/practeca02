# word_exporter.py
from docx import Document
from datetime import datetime

def save_to_word(data, title, filename=None):
    """Сохранение данных в Word"""
    doc = Document()
    doc.add_heading(title, 0).alignment = 1
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph()
    
    if data and len(data) > 0:
        table = doc.add_table(rows=1, cols=len(data[0]) + 1)
        table.style = 'Light Grid Accent 1'
        
        # Заголовки
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        headers = ['Название', 'Производитель', 'Количество', 'Дата', 'Цена', 'Причина']
        for i in range(len(data[0])):
            if i < len(headers):
                hdr_cells[i + 1].text = headers[i]
        
        # Данные
        for i, row in enumerate(data, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            for j, val in enumerate(row):
                if j + 1 < len(row_cells):
                    if hasattr(val, 'strftime'):
                        row_cells[j + 1].text = val.strftime('%d.%m.%Y')
                    else:
                        row_cells[j + 1].text = str(val)
    
    if not filename:
        filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    doc.save(filename)
    return filename